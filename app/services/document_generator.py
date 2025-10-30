import openai
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from config import Config

openai.api_key = Config.OPENAI_API_KEY


class DocumentGenerator:

    """
    Service for generating AI-powered business documents
    """
    
    def __init__(self):
        self.upload_folder = Config.UPLOAD_FOLDER
        os.makedirs(self.upload_folder, exist_ok=True)
    
    def generate_document(self, document_type, submission_data, custom_prompt=''):
        """
        Generate a document based on type and submission data
        """
        try:
            if document_type == 'pitch_deck':
                return self._generate_pitch_deck(submission_data, custom_prompt)
            elif document_type == 'business_plan':
                return self._generate_business_plan(submission_data, custom_prompt)
            elif document_type == 'financial_model':
                return self._generate_financial_model(submission_data, custom_prompt)
            elif document_type == 'executive_summary':
                return self._generate_executive_summary(submission_data, custom_prompt)
            else:
                return {'success': False, 'error': 'Invalid document type'}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _generate_ai_content(self, prompt, max_tokens=2000):
        """
        Generate content using OpenAI API
        """
        try:
            response = openai.ChatCompletion.create(
                model=Config.OPENAI_MODEL,
                messages=[
                    {"role": "system", "content": "You are an expert business consultant and document writer specializing in startup pitch decks, business plans, and financial models."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=max_tokens,
                temperature=0.7
            )
            
            return {
                'success': True,
                'content': response.choices[0].message.content,
                'tokens': response.usage.total_tokens
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _generate_pitch_deck(self, submission_data, custom_prompt):
        """
        Generate a pitch deck document
        """
        try:
            startup_name = submission_data.get('startupName', 'Startup')
            
            # Generate AI content for each slide
            prompt = f"""
            Create a comprehensive pitch deck content for the following startup:
            
            Startup Name: {startup_name}
            Industry: {submission_data.get('industry', 'N/A')}
            Stage: {submission_data.get('stage', 'N/A')}
            Description: {submission_data.get('description', 'N/A')}
            Problem: {submission_data.get('problemStatement', 'N/A')}
            Solution: {submission_data.get('solution', 'N/A')}
            Target Market: {submission_data.get('targetMarket', 'N/A')}
            Business Model: {submission_data.get('businessModel', 'N/A')}
            Competition: {submission_data.get('competition', 'N/A')}
            Team Size: {submission_data.get('teamSize', 'N/A')}
            Funding Required: ${submission_data.get('fundingRequired', 0)}
            
            {custom_prompt}
            
            Generate content for the following slides:
            1. Title Slide
            2. Problem Statement
            3. Solution
            4. Market Opportunity
            5. Business Model
            6. Competition & Competitive Advantage
            7. Go-to-Market Strategy
            8. Financial Projections
            9. Team
            10. Ask & Use of Funds
            
            Format each slide with a clear heading and bullet points.
            """
            
            ai_result = self._generate_ai_content(prompt, max_tokens=3000)
            
            if not ai_result['success']:
                return ai_result
            
            # Create Word document
            doc = Document()
            
            # Title
            title = doc.add_heading(f'{startup_name} - Pitch Deck', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add generated content
            content_lines = ai_result['content'].split('\n')
            for line in content_lines:
                if line.strip():
                    if line.startswith('#') or line.endswith(':'):
                        doc.add_heading(line.strip('#').strip(':').strip(), level=1)
                    elif line.startswith('-') or line.startswith('•'):
                        doc.add_paragraph(line.strip('-').strip('•').strip(), style='List Bullet')
                    else:
                        doc.add_paragraph(line.strip())
            
            # Add footer
            doc.add_paragraph()
            footer = doc.add_paragraph(f'Generated on {datetime.now().strftime("%B %d, %Y")}')
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
            filename = f'{startup_name.replace(" ", "_")}_PitchDeck_{timestamp}.docx'
            filepath = os.path.join(self.upload_folder, filename)
            doc.save(filepath)
            
            file_size = os.path.getsize(filepath)
            
            return {
                'success': True,
                'title': f'{startup_name} - Pitch Deck',
                'file_name': filename,
                'file_path': filepath,
                'file_size': file_size,
                'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'tokens': ai_result['tokens']
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _generate_business_plan(self, submission_data, custom_prompt):
        """
        Generate a comprehensive business plan
        """
        try:
            startup_name = submission_data.get('startupName', 'Startup')
            
            prompt = f"""
            Create a detailed business plan for the following startup:
            
            Startup Name: {startup_name}
            Industry: {submission_data.get('industry', 'N/A')}
            Stage: {submission_data.get('stage', 'N/A')}
            Description: {submission_data.get('description', 'N/A')}
            Problem: {submission_data.get('problemStatement', 'N/A')}
            Solution: {submission_data.get('solution', 'N/A')}
            Target Market: {submission_data.get('targetMarket', 'N/A')}
            Market Size: {submission_data.get('marketSize', 'N/A')}
            Business Model: {submission_data.get('businessModel', 'N/A')}
            Revenue Streams: {submission_data.get('revenueStreams', 'N/A')}
            Competition: {submission_data.get('competition', 'N/A')}
            Competitive Advantage: {submission_data.get('competitiveAdvantage', 'N/A')}
            Team: {submission_data.get('teamDescription', 'N/A')}
            Current Traction: {submission_data.get('currentTraction', 'N/A')}
            Funding Required: ${submission_data.get('fundingRequired', 0)}
            Current Revenue: ${submission_data.get('currentRevenue', 0)}
            
            {custom_prompt}
            
            Generate a comprehensive business plan with the following sections:
            1. Executive Summary
            2. Company Description
            3. Market Analysis
            4. Organization and Management
            5. Service or Product Line
            6. Marketing and Sales Strategy
            7. Financial Projections
            8. Funding Requirements
            9. Appendix
            
            Make it detailed, professional, and investor-ready.
            """
            
            ai_result = self._generate_ai_content(prompt, max_tokens=4000)
            
            if not ai_result['success']:
                return ai_result
            
            # Create Word document
            doc = Document()
            
            # Title page
            title = doc.add_heading(f'{startup_name}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading('Business Plan', level=1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()
            
            # Table of Contents
            doc.add_heading('Table of Contents', level=1)
            toc_items = [
                'Executive Summary',
                'Company Description',
                'Market Analysis',
                'Organization and Management',
                'Service or Product Line',
                'Marketing and Sales Strategy',
                'Financial Projections',
                'Funding Requirements',
                'Appendix'
            ]
            for i, item in enumerate(toc_items, 1):
                doc.add_paragraph(f'{i}. {item}', style='List Number')
            
            doc.add_page_break()
            
            # Add generated content
            content_lines = ai_result['content'].split('\n')
            for line in content_lines:
                if line.strip():
                    if line.startswith('#') or (line.isupper() and len(line.split()) <= 5):
                        doc.add_heading(line.strip('#').strip(), level=1)
                    elif line.startswith('##'):
                        doc.add_heading(line.strip('#').strip(), level=2)
                    elif line.startswith('-') or line.startswith('•'):
                        doc.add_paragraph(line.strip('-').strip('•').strip(), style='List Bullet')
                    else:
                        doc.add_paragraph(line.strip())
            
            # Save document
            timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
            filename = f'{startup_name.replace(" ", "_")}_BusinessPlan_{timestamp}.docx'
            filepath = os.path.join(self.upload_folder, filename)
            doc.save(filepath)
            
            file_size = os.path.getsize(filepath)
            
            return {
                'success': True,
                'title': f'{startup_name} - Business Plan',
                'file_name': filename,
                'file_path': filepath,
                'file_size': file_size,
                'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'tokens': ai_result['tokens']
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _generate_financial_model(self, submission_data, custom_prompt):
        """
        Generate a financial model document
        """
        try:
            startup_name = submission_data.get('startupName', 'Startup')
            
            prompt = f"""
            Create a detailed 3-year financial projection for the following startup:
            
            Startup Name: {startup_name}
            Industry: {submission_data.get('industry', 'N/A')}
            Business Model: {submission_data.get('businessModel', 'N/A')}
            Revenue Streams: {submission_data.get('revenueStreams', 'N/A')}
            Current Revenue: ${submission_data.get('currentRevenue', 0)}
            Funding Required: ${submission_data.get('fundingRequired', 0)}
            Monthly Burn Rate: ${submission_data.get('monthlyBurnRate', 0)}
            Team Size: {submission_data.get('teamSize', 1)}
            
            {custom_prompt}
            
            Generate a comprehensive financial model including:
            1. Revenue Projections (Year 1, 2, 3)
            2. Cost Structure
            3. Profit & Loss Statement
            4. Cash Flow Projection
            5. Break-even Analysis

            6. Key Financial Metrics (CAC, LTV, Gross Margin, etc.)
            7. Funding Allocation
            8. Growth Assumptions
            
            Provide realistic numbers and clear explanations for each projection.
            """
            
            ai_result = self._generate_ai_content(prompt, max_tokens=3500)
            
            if not ai_result['success']:
                return ai_result
            
            # Create Word document
            doc = Document()
            
            # Title
            title = doc.add_heading(f'{startup_name} - Financial Model', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading('3-Year Financial Projections', level=1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()
            
            # Add generated content
            content_lines = ai_result['content'].split('\n')
            for line in content_lines:
                if line.strip():
                    if line.startswith('#') or (line.isupper() and len(line.split()) <= 6):
                        doc.add_heading(line.strip('#').strip(), level=1)
                    elif line.startswith('##'):
                        doc.add_heading(line.strip('#').strip(), level=2)
                    elif line.startswith('-') or line.startswith('•'):
                        doc.add_paragraph(line.strip('-').strip('•').strip(), style='List Bullet')
                    else:
                        doc.add_paragraph(line.strip())
            
            # Add disclaimer
            doc.add_page_break()
            doc.add_heading('Disclaimer', level=1)
            disclaimer_text = """
            This financial model is based on assumptions and projections. Actual results may vary significantly. 
            This document should not be considered as financial advice. Please consult with financial professionals 
            before making any investment decisions.
            """
            doc.add_paragraph(disclaimer_text)
            
            # Save document
            timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
            filename = f'{startup_name.replace(" ", "_")}_FinancialModel_{timestamp}.docx'
            filepath = os.path.join(self.upload_folder, filename)
            doc.save(filepath)
            
            file_size = os.path.getsize(filepath)
            
            return {
                'success': True,
                'title': f'{startup_name} - Financial Model',
                'file_name': filename,
                'file_path': filepath,
                'file_size': file_size,
                'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'tokens': ai_result['tokens']
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _generate_executive_summary(self, submission_data, custom_prompt):
        """
        Generate an executive summary
        """
        try:
            startup_name = submission_data.get('startupName', 'Startup')
            
            prompt = f"""
            Create a compelling executive summary for the following startup:
            
            Startup Name: {startup_name}
            Industry: {submission_data.get('industry', 'N/A')}
            Stage: {submission_data.get('stage', 'N/A')}
            Description: {submission_data.get('description', 'N/A')}
            Problem: {submission_data.get('problemStatement', 'N/A')}
            Solution: {submission_data.get('solution', 'N/A')}
            Unique Value Proposition: {submission_data.get('uniqueValueProposition', 'N/A')}
            Target Market: {submission_data.get('targetMarket', 'N/A')}
            Market Size: {submission_data.get('marketSize', 'N/A')}
            Business Model: {submission_data.get('businessModel', 'N/A')}
            Competitive Advantage: {submission_data.get('competitiveAdvantage', 'N/A')}
            Current Traction: {submission_data.get('currentTraction', 'N/A')}
            Team Size: {submission_data.get('teamSize', 'N/A')}
            Funding Required: ${submission_data.get('fundingRequired', 0)}
            Current Revenue: ${submission_data.get('currentRevenue', 0)}
            
            {custom_prompt}
            
            Create a concise, compelling executive summary (2-3 pages) that includes:
            1. Company Overview
            2. Problem & Solution
            3. Market Opportunity
            4. Business Model
            5. Competitive Advantage
            6. Traction & Milestones
            7. Financial Highlights
            8. Funding Ask & Use of Funds
            9. Vision & Next Steps
            
            Make it investor-ready and compelling.
            """
            
            ai_result = self._generate_ai_content(prompt, max_tokens=2500)
            
            if not ai_result['success']:
                return ai_result
            
            # Create Word document
            doc = Document()
            
            # Title
            title = doc.add_heading(f'{startup_name}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading('Executive Summary', level=1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()
            
            # Add generated content
            content_lines = ai_result['content'].split('\n')
            for line in content_lines:
                if line.strip():
                    if line.startswith('#'):
                        doc.add_heading(line.strip('#').strip(), level=1)
                    elif line.startswith('##'):
                        doc.add_heading(line.strip('#').strip(), level=2)
                    elif line.startswith('-') or line.startswith('•'):
                        doc.add_paragraph(line.strip('-').strip('•').strip(), style='List Bullet')
                    else:
                        doc.add_paragraph(line.strip())
            
            # Save document
            timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
            filename = f'{startup_name.replace(" ", "_")}_ExecutiveSummary_{timestamp}.docx'
            filepath = os.path.join(self.upload_folder, filename)
            doc.save(filepath)
            
            file_size = os.path.getsize(filepath)
            
            return {
                'success': True,
                'title': f'{startup_name} - Executive Summary',
                'file_name': filename,
                'file_path': filepath,
                'file_size': file_size,
                'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'tokens': ai_result['tokens']
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}